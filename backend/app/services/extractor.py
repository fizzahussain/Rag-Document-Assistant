import io
import json
import shutil
import subprocess
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

import docx
import fitz
import pandas as pd
from bs4 import BeautifulSoup
from pydantic import BaseModel

from backend.app.config import settings
from backend.app.core.exceptions import ExtractionError, OCRRequiredError


class ExtractedPage(BaseModel):
    """Represent text extracted from an individual page"""

    page_number: int
    text: str


class ExtractedDocument(BaseModel):
    """Contain document extraction output"""

    text: str
    pages: list[ExtractedPage]
    metadata: dict[str, Any]


class BaseExtractor(ABC):
    """Define the interface for file text extractors"""

    @abstractmethod
    def extract(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        """Extract text and metadata from file content"""


class PDFExtractor(BaseExtractor):
    """Extract PDF text and selectively OCR pages that need it"""

    @staticmethod
    def _clean_char_count(text: str) -> int:
        return len("".join(text.split()))

    @staticmethod
    def _tesseract_cli_ocr(page: fitz.Page) -> str:
        tesseract = shutil.which("tesseract")
        if not tesseract:
            return ""

        scale = settings.OCR_DPI / 72
        pixmap = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
        png_bytes = pixmap.tobytes("png")
        try:
            result = subprocess.run(
                [
                    tesseract,
                    "stdin",
                    "stdout",
                    "-l",
                    settings.OCR_LANGUAGE,
                    "--dpi",
                    str(settings.OCR_DPI),
                ],
                input=png_bytes,
                capture_output=True,
                check=False,
                timeout=90,
            )
        except (OSError, subprocess.SubprocessError):
            return ""
        if result.returncode != 0:
            return ""
        return result.stdout.decode("utf-8", errors="replace").strip()

    def _extract_page_text(self, page: fitz.Page, page_number: int) -> tuple[str, str]:
        native_text = page.get_text("text").strip()
        if self._clean_char_count(native_text) >= settings.OCR_MIN_TEXT_CHARS:
            return native_text, "native"

        if not settings.OCR_ENABLED:
            return native_text, "native"

        ocr_text = ""
        try:
            text_page = page.get_textpage_ocr(
                language=settings.OCR_LANGUAGE,
                dpi=settings.OCR_DPI,
                full=True,
            )
            ocr_text = page.get_text("text", textpage=text_page).strip()
        except Exception:
            ocr_text = ""

        if self._clean_char_count(ocr_text) < settings.OCR_MIN_TEXT_CHARS:
            cli_text = self._tesseract_cli_ocr(page)
            if self._clean_char_count(cli_text) > self._clean_char_count(ocr_text):
                ocr_text = cli_text

        if self._clean_char_count(ocr_text) > self._clean_char_count(native_text):
            return ocr_text, "ocr"
        if native_text:
            return native_text, "native"

        raise OCRRequiredError(
            f"Page {page_number} requires OCR, but Tesseract produced no readable text. "
            "Verify that Tesseract and the configured language data are available."
        )

    def extract(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        document = None

        try:
            document = fitz.open(stream=file_bytes, filetype="pdf")
            pages: list[ExtractedPage] = []
            full_text_list: list[str] = []
            ocr_pages: list[int] = []
            native_pages: list[int] = []

            for page_index, page in enumerate(document):
                page_number = page_index + 1
                page_text, method = self._extract_page_text(page, page_number)
                pages.append(
                    ExtractedPage(
                        page_number=page_number,
                        text=page_text,
                    )
                )

                if page_text:
                    full_text_list.append(page_text)

                if method == "ocr":
                    ocr_pages.append(page_number)
                else:
                    native_pages.append(page_number)

            full_text = "\n\n".join(full_text_list).strip()
            total_clean_chars = self._clean_char_count(full_text)

            if total_clean_chars < settings.OCR_MIN_TEXT_CHARS:
                raise OCRRequiredError(
                    f"PDF document '{filename}' contains little or no readable text "
                    f"after extraction ({total_clean_chars} chars)."
                )

            metadata = {
                "page_count": len(document),
                "author": document.metadata.get("author", ""),
                "title": document.metadata.get("title", ""),
                "format": "PDF",
                "ocr_pages": ocr_pages,
                "native_text_pages": native_pages,
                "ocr_enabled": settings.OCR_ENABLED,
            }

            return ExtractedDocument(
                text=full_text,
                pages=pages,
                metadata=metadata,
            )
        except OCRRequiredError:
            raise
        except Exception as exc:
            raise ExtractionError(
                f"Failed to extract PDF content from '{filename}': {exc!s}"
            ) from exc
        finally:
            if document is not None:
                document.close()


class DOCXExtractor(BaseExtractor):
    """Extract paragraphs and table rows from DOCX files"""

    def extract(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        try:
            document = docx.Document(io.BytesIO(file_bytes))
            paragraphs: list[str] = []

            for paragraph in document.paragraphs:
                text = paragraph.text.strip()

                if text:
                    paragraphs.append(text)

            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )

                    if row_text:
                        paragraphs.append(row_text)

            full_text = "\n\n".join(paragraphs)
            pages = [
                ExtractedPage(
                    page_number=1,
                    text=full_text,
                )
            ]
            metadata = {
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "format": "DOCX",
            }

            return ExtractedDocument(
                text=full_text,
                pages=pages,
                metadata=metadata,
            )
        except Exception as exc:
            raise ExtractionError(
                f"Failed to extract DOCX content from '{filename}': {exc!s}"
            ) from exc


class TextExtractor(BaseExtractor):
    """Extract plain text from TXT and Markdown files"""

    def extract(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        try:
            try:
                full_text = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                full_text = file_bytes.decode("latin-1")

            full_text = full_text.strip()
            pages = [
                ExtractedPage(
                    page_number=1,
                    text=full_text,
                )
            ]
            metadata = {"format": Path(filename).suffix.upper().lstrip(".")}

            return ExtractedDocument(
                text=full_text,
                pages=pages,
                metadata=metadata,
            )
        except Exception as exc:
            raise ExtractionError(
                f"Failed to extract text content from '{filename}': {exc!s}"
            ) from exc


class CSVExtractor(BaseExtractor):
    """Extract structured text from CSV files"""

    def extract(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        try:
            try:
                content = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                content = file_bytes.decode("latin-1")

            dataframe = pd.read_csv(io.StringIO(content))
            lines: list[str] = []

            headers = [str(column) for column in dataframe.columns]
            lines.append("Columns: " + ", ".join(headers))

            for row_index, row in dataframe.iterrows():
                row_text = " | ".join(
                    f"{column}: {value}" for column, value in row.items() if pd.notna(value)
                )
                lines.append(f"Row {row_index + 1}: {row_text}")

            full_text = "\n".join(lines)
            pages = [
                ExtractedPage(
                    page_number=1,
                    text=full_text,
                )
            ]
            metadata = {
                "row_count": len(dataframe),
                "column_count": len(dataframe.columns),
                "format": "CSV",
            }

            return ExtractedDocument(
                text=full_text,
                pages=pages,
                metadata=metadata,
            )
        except Exception as exc:
            raise ExtractionError(
                f"Failed to extract CSV content from '{filename}': {exc!s}"
            ) from exc


class HTMLExtractor(BaseExtractor):
    """Extract visible text from HTML documents"""

    def extract(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        try:
            try:
                html = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                html = file_bytes.decode("latin-1")

            soup = BeautifulSoup(html, "html.parser")

            removable_tags = [
                "script",
                "style",
                "meta",
                "noscript",
                "header",
                "footer",
            ]

            for tag in soup(removable_tags):
                tag.decompose()

            full_text = soup.get_text(separator="\n", strip=True)
            pages = [
                ExtractedPage(
                    page_number=1,
                    text=full_text,
                )
            ]
            metadata = {
                "title": soup.title.string if soup.title else "",
                "format": "HTML",
            }

            return ExtractedDocument(
                text=full_text,
                pages=pages,
                metadata=metadata,
            )
        except Exception as exc:
            raise ExtractionError(
                f"Failed to extract HTML content from '{filename}': {exc!s}"
            ) from exc


class JSONExtractor(BaseExtractor):
    """Extract formatted text from JSON files"""

    def extract(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        try:
            try:
                json_text = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                json_text = file_bytes.decode("latin-1")

            data = json.loads(json_text)
            full_text = json.dumps(data, indent=2)
            pages = [
                ExtractedPage(
                    page_number=1,
                    text=full_text,
                )
            ]
            metadata = {"format": "JSON"}

            return ExtractedDocument(
                text=full_text,
                pages=pages,
                metadata=metadata,
            )
        except Exception as exc:
            raise ExtractionError(
                f"Failed to extract JSON content from '{filename}': {exc!s}"
            ) from exc


class ExtractorFactory:
    """Return the appropriate extractor for a filename"""

    @staticmethod
    def get_extractor(filename: str) -> BaseExtractor:
        extension = Path(filename).suffix.lower().lstrip(".")

        extractors: dict[str, type[BaseExtractor]] = {
            "pdf": PDFExtractor,
            "docx": DOCXExtractor,
            "txt": TextExtractor,
            "md": TextExtractor,
            "markdown": TextExtractor,
            "csv": CSVExtractor,
            "html": HTMLExtractor,
            "htm": HTMLExtractor,
            "json": JSONExtractor,
        }

        extractor_class = extractors.get(extension)

        if extractor_class is None:
            raise ExtractionError(f"No extractor registered for file extension '.{extension}'")

        return extractor_class()
